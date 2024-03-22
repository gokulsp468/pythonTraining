from docx import Document
from docx.shared import Inches
from docx.shared import Pt,RGBColor

from docx.enum.text import WD_UNDERLINE
from fpdf import FPDF
import os



records = [
    {'name': "sachinTendulkar", 'description': 'Sachin Ramesh Tendulkar, is an Indian former international cricketer who captained the Indian national team. ', 'image': 'Sachin Tendulkar.jpeg','runs':50},

      {'name': "Dhoni", 'description': 'MahendraSinghDhoni is an Indian professional cricketer. He is a right handed batter and a wicket-keeper.', 'image': '_.jpeg','runs':40},

       {'name': "sanju", 'description': 'SanjuViswanath Samson is an Indian international cricketer, who plays for Kerala in domestic cricket and Rajasthan Royals in the Indian Premier League.', 'image': 'sanju.jpeg','runs':45},

         {'name': "Chris Gayle", 'description':' Christopher Henry Gayle OD is a Jamaican cricketer who has played international cricket for the West Indies from 1999 to 2021. Nicknamed "The Universe Boss"', 'image': '_ (1).jpeg','runs':145},
    
    
]
document = Document()

heading_paragraph = document.add_paragraph('SCORE BOARD')
heading_paragraph.alignment = 1

heading_run = heading_paragraph.runs[0]
heading_run.font.size = Pt(20)
heading_run.font.color.rgb = RGBColor(0,0,255)
heading_run.font.underline = WD_UNDERLINE.SINGLE 


table = document.add_table(rows=1, cols=4)
header_cells = table.rows[0].cells
header_cells[0].text = 'NAME'
header_cells[1].text = 'DESCRIPTION'
header_cells[2].text = 'IMAGE'
header_cells[3].text = 'SCORE'



for record in records:
   
    
    row_cells = table.add_row().cells
    row_cells[0].text = record['name']
    row_cells[1].text=record['description']
    row_cells[2].paragraphs[0].add_run().add_picture(record['image'], width=Inches(1.25),height=Inches(1.5))   
    row_cells[3].text = str(record['runs'])
table.style = 'Table Grid'
document.add_page_break()

document.save('demo.docx')

word_file = 'demo.docx'
document.save(word_file)
pdf_file = 'Table_scores.pdf'
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
docx_to_pdf_command = f'libreoffice --headless --convert-to pdf {word_file}'


os.system(docx_to_pdf_command)




